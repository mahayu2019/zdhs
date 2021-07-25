#!/usr/bin/env python
# coding=utf-8

'''
批量生成word合同
注意:excel表的表头必须和word中的替换标记字符一致
'''

from openpyxl import load_workbook
from docx import Document


# 读取和替换文本
def info_update(doc, old_info, new_info):
    for para in doc.paragraphs:  # paragraphs为文档段落集合,paragraph代表一个段落
        for run in para.runs:  # run为具有相同样式的一段连续文本,runs为run的集合
            run.text = run.text.replace(old_info, new_info)
    # 以下代码适用表格类型的替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info)


wb = load_workbook('合同信息.xlsx')  # 打开填充内容的工作簿
ws = wb.active  # 激活工作表

for row in range(2, ws.max_row + 1):  # 2-->跳过标题行,ws.max_row+1包含到最后一行
    doc = Document('合同模板.docx')  # 打开模板文档
    for col in range(1, ws.max_column + 1):  # 遍历列数据,同理最大值+1
        old_info = str(ws.cell(row=1, column=col).value)
        new_info = str(ws.cell(row=row, column=col).value)
        info_update(doc, old_info, new_info)  # 调用替换函数替换值
    com_name = str(ws.cell(row=row, column=2).value)
    doc.save(f'{com_name}合同.docx')
